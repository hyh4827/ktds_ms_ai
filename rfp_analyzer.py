import json
import os
from datetime import datetime
from typing import Dict, List, Any
import pandas as pd
import numpy as np
from dotenv import load_dotenv
from azure.search.documents import SearchClient
from azure.search.documents.indexes import SearchIndexClient
from azure.search.documents.models import VectorizedQuery
from azure.core.credentials import AzureKeyCredential
from azure.search.documents.indexes.models import (
    SearchIndex,
    SearchField,
    SearchFieldDataType,
    SimpleField,
    SearchableField,
    VectorSearch,
    HnswAlgorithmConfiguration,
    VectorSearchProfile,
    SemanticConfiguration,
    SemanticPrioritizedFields,
    SemanticField,
    SemanticSearch
)
import openai
from openai import AzureOpenAI
import PyPDF2
import docx
from io import BytesIO
import re
import streamlit as st

class RFPAnalyzer:
    def __init__(self):
        self.search_client = None
        self.openai_client = None
        self.index_name = "rfp-documents"
        
    def initialize_services(self, search_endpoint, search_key, openai_endpoint, openai_key, openai_api_version):
        """Azure 서비스 초기화"""
        try:
            # Azure AI Search 클라이언트 초기화
            credential = AzureKeyCredential(search_key)
            self.search_client = SearchClient(
                endpoint=search_endpoint,
                index_name=self.index_name,
                credential=credential
            )
            
            # Azure OpenAI 클라이언트 초기화
            self.openai_client = AzureOpenAI(
                azure_endpoint=openai_endpoint,
                api_key=openai_key,
                api_version=openai_api_version
            )
            
            return True
        except Exception as e:
            st.error(f"❌ 서비스 초기화 오류: {str(e)}")
            return False
    
    def create_search_index(self, search_endpoint, search_key):
        """Azure AI Search 인덱스 생성"""
        try:
            credential = AzureKeyCredential(search_key)
            index_client = SearchIndexClient(endpoint=search_endpoint, credential=credential)
            
            # 인덱스 정의
            fields = [
                SimpleField(name="id", type=SearchFieldDataType.String, key=True),
                SearchableField(name="title", type=SearchFieldDataType.String),
                SearchableField(name="content", type=SearchFieldDataType.String),
                SearchableField(name="requirements", type=SearchFieldDataType.String),
                SearchableField(name="project_type", type=SearchFieldDataType.String),
                SearchableField(name="budget_range", type=SearchFieldDataType.String),
                SimpleField(name="submission_deadline", type=SearchFieldDataType.String),
                SearchableField(name="evaluation_criteria", type=SearchFieldDataType.String),
                SimpleField(name="created_date", type=SearchFieldDataType.DateTimeOffset),
                SearchField(name="content_vector", type=SearchFieldDataType.Collection(SearchFieldDataType.Single),
                           searchable=True, vector_search_dimensions=1536, vector_search_profile_name="myHnswProfile")
            ]
            
            # Vector search 구성
            vector_search = VectorSearch(
                algorithms=[
                    HnswAlgorithmConfiguration(name="myHnsw")
                ],
                profiles=[
                    VectorSearchProfile(
                        name="myHnswProfile",
                        algorithm_configuration_name="myHnsw"
                    )
                ]
            )
            
            # Semantic search 구성
            semantic_config = SemanticConfiguration(
                name="my-semantic-config",
                prioritized_fields=SemanticPrioritizedFields(
                    title_field=SemanticField(field_name="title"),
                    content_fields=[SemanticField(field_name="content")]
                )
            )
            
            semantic_search = SemanticSearch(configurations=[semantic_config])
            
            # 인덱스 생성
            index = SearchIndex(
                name=self.index_name,
                fields=fields,
                vector_search=vector_search,
                semantic_search=semantic_search
            )
            
            index_client.create_or_update_index(index)
            return True
            
        except Exception as e:
            st.error(f"❌ 인덱스 생성 오류: {str(e)}")
            return False
    
    def get_embedding(self, text: str) -> List[float]:
        """텍스트의 임베딩 벡터 생성"""
        try:
            # 텍스트가 너무 길 경우 청크로 나누어 처리
            max_chunk_size = 4000  # 토큰 제한을 고려하여 4000자로 설정 (약 6000 토큰)
            
            if len(text) <= max_chunk_size:
                # 텍스트가 짧으면 그대로 처리
                response = self.openai_client.embeddings.create(
                    input=text,
                    model="text-embedding-ada-002"
                )
                return response.data[0].embedding
            else:
                # 텍스트가 길면 청크로 나누어 처리
                st.warning(f"⚠️ 텍스트가 길어서 청크 단위로 처리합니다. (길이: {len(text)} 문자)")
                
                # 텍스트를 청크로 분할
                chunks = self._split_text_into_chunks(text, max_chunk_size)
                embeddings = []
                
                progress_bar = st.progress(0)
                for i, chunk in enumerate(chunks):
                    st.write(f"   청크 {i+1}/{len(chunks)} 처리 중...")
                    response = self.openai_client.embeddings.create(
                        input=chunk,
                        model="text-embedding-ada-002"
                    )
                    embeddings.append(response.data[0].embedding)
                    progress_bar.progress((i + 1) / len(chunks))
                
                # 모든 청크의 임베딩을 평균화하여 하나의 벡터로 만듦
                if embeddings:
                    avg_embedding = np.mean(embeddings, axis=0).tolist()
                    return avg_embedding
                else:
                    return []
                    
        except Exception as e:
            st.error(f"❌ 임베딩 생성 오류: {str(e)}")
            return []
    
    def _split_text_into_chunks(self, text: str, max_chunk_size: int) -> List[str]:
        """텍스트를 청크로 분할하는 함수"""
        chunks = []
        
        # 문장 단위로 분할을 시도
        sentences = text.split('. ')
        current_chunk = ""
        
        for sentence in sentences:
            # 현재 청크에 문장을 추가했을 때의 길이 확인
            test_chunk = current_chunk + sentence + ". " if current_chunk else sentence + ". "
            
            if len(test_chunk) <= max_chunk_size:
                current_chunk = test_chunk
            else:
                # 현재 청크가 비어있지 않으면 저장
                if current_chunk:
                    chunks.append(current_chunk.strip())
                    current_chunk = sentence + ". "
                else:
                    # 문장 자체가 너무 긴 경우 강제로 자름
                    if len(sentence) > max_chunk_size:
                        # 문장을 단어 단위로 자름
                        words = sentence.split()
                        temp_chunk = ""
                        
                        for word in words:
                            if len(temp_chunk + " " + word) <= max_chunk_size:
                                temp_chunk += " " + word if temp_chunk else word
                            else:
                                if temp_chunk:
                                    chunks.append(temp_chunk.strip())
                                temp_chunk = word
                        
                        if temp_chunk:
                            current_chunk = temp_chunk + ". "
                    else:
                        current_chunk = sentence + ". "
        
        # 마지막 청크 추가
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        # 각 청크의 크기를 다시 확인하고 필요시 추가 분할
        final_chunks = []
        for chunk in chunks:
            if len(chunk) <= max_chunk_size:
                final_chunks.append(chunk)
            else:
                # 청크가 여전히 크면 추가로 분할
                words = chunk.split()
                temp_chunk = ""
                
                for word in words:
                    if len(temp_chunk + " " + word) <= max_chunk_size:
                        temp_chunk += " " + word if temp_chunk else word
                    else:
                        if temp_chunk:
                            final_chunks.append(temp_chunk.strip())
                        temp_chunk = word
                
                if temp_chunk:
                    final_chunks.append(temp_chunk.strip())
        
        return final_chunks
    
    def extract_text_from_file(self, file_path):
        """파일에서 텍스트 추출"""
        try:
            # 파일 경로 정규화 및 존재 확인
            file_path = os.path.abspath(file_path)
            st.info(f"🔍 파일 경로: {file_path}")
            
            if not os.path.exists(file_path):
                st.error(f"❌ 파일이 존재하지 않습니다: {file_path}")
                return None, None
            
            if not os.path.isfile(file_path):
                st.error(f"❌ 경로가 파일이 아닙니다: {file_path}")
                return None, None
            
            # 파일 크기 확인
            file_size = os.path.getsize(file_path)
            st.info(f"📊 파일 크기: {file_size} bytes")
            
            if file_size == 0:
                st.error("❌ 파일이 비어있습니다.")
                return None, None
            
            file_extension = os.path.splitext(file_path)[1].lower()
            st.info(f"📄 파일 확장자: {file_extension}")
            
            if file_extension == ".pdf":
                st.info("📖 PDF 파일 처리 중...")
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    total_pages = len(pdf_reader.pages)
                    st.info(f"📄 PDF 페이지 수: {total_pages}")
                    
                    # PDF 제목 추출 시도
                    pdf_title = None
                    try:
                        if pdf_reader.metadata and pdf_reader.metadata.title:
                            pdf_title = pdf_reader.metadata.title.strip()
                            st.info(f"📋 PDF 제목: {pdf_title}")
                    except:
                        pass
                    
                    # 제목이 없으면 파일명 사용
                    if not pdf_title:
                        pdf_title = os.path.splitext(os.path.basename(file_path))[0]
                        st.info(f"📋 파일명을 제목으로 사용: {pdf_title}")
                    
                    # 페이지 수 제한 (200페이지)
                    max_pages = 200
                    if total_pages > max_pages:
                        st.warning(f"⚠️ PDF 페이지 수가 {max_pages}페이지를 초과합니다. 처음 {max_pages}페이지만 처리합니다.")
                        pages_to_process = max_pages
                    else:
                        pages_to_process = total_pages
                    
                    text = ""
                    progress_bar = st.progress(0)
                    for i in range(pages_to_process):
                        try:
                            page = pdf_reader.pages[i]
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                            # st.write(f"   페이지 {i+1} 처리 완료")
                            progress_bar.progress((i + 1) / pages_to_process)
                        except Exception as page_error:
                            st.warning(f"   페이지 {i+1} 처리 오류: {str(page_error)}")
                            continue
                    
                    if not text.strip():
                        st.error("❌ PDF에서 텍스트를 추출할 수 없습니다.")
                        return None, None
                    
                    st.success(f"✅ PDF 텍스트 추출 완료: {len(text)} 문자 (처리된 페이지: {pages_to_process}/{total_pages})")
                    return text, pdf_title
            
            elif file_extension == ".docx":
                st.info("📝 DOCX 파일 처리 중...")
                doc = docx.Document(file_path)
                st.info(f"📄 DOCX 단락 수: {len(doc.paragraphs)}")
                
                # DOCX 제목 추출 시도 (첫 번째 단락을 제목으로 사용)
                docx_title = None
                if doc.paragraphs:
                    first_paragraph = doc.paragraphs[0].text.strip()
                    if first_paragraph and len(first_paragraph) < 100:  # 너무 길면 제목이 아닐 가능성
                        docx_title = first_paragraph
                        st.info(f"📋 DOCX 제목: {docx_title}")
                
                # 제목이 없으면 파일명 사용
                if not docx_title:
                    docx_title = os.path.splitext(os.path.basename(file_path))[0]
                    st.info(f"📋 파일명을 제목으로 사용: {docx_title}")
                
                text = ""
                for i, paragraph in enumerate(doc.paragraphs):
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"
                
                if not text.strip():
                    st.error("❌ DOCX에서 텍스트를 추출할 수 없습니다.")
                    return None, None
                
                st.success(f"✅ DOCX 텍스트 추출 완료: {len(text)} 문자")
                return text, docx_title
            
            elif file_extension == ".txt":
                st.info("📄 TXT 파일 처리 중...")
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()
                    if not content.strip():
                        st.error("❌ TXT 파일이 비어있습니다.")
                        return None, None
                    
                    # TXT 제목 추출 시도 (첫 번째 줄을 제목으로 사용)
                    txt_title = None
                    first_line = content.split('\n')[0].strip()
                    if first_line and len(first_line) < 100:  # 너무 길면 제목이 아닐 가능성
                        txt_title = first_line
                        st.info(f"📋 TXT 제목: {txt_title}")
                    
                    # 제목이 없으면 파일명 사용
                    if not txt_title:
                        txt_title = os.path.splitext(os.path.basename(file_path))[0]
                        st.info(f"📋 파일명을 제목으로 사용: {txt_title}")
                    
                    st.success(f"✅ TXT 텍스트 추출 완료: {len(content)} 문자")
                    return content, txt_title
            
            else:
                st.error(f"❌ 지원하지 않는 파일 형식입니다: {file_extension}")
                st.info("   지원 형식: .pdf, .docx, .txt")
                return None, None
                
        except FileNotFoundError as e:
            st.error(f"❌ 파일을 찾을 수 없습니다: {str(e)}")
            return None, None
        except PermissionError as e:
            st.error(f"❌ 파일 접근 권한이 없습니다: {str(e)}")
            return None, None
        except UnicodeDecodeError as e:
            st.error(f"❌ 파일 인코딩 오류: {str(e)}")
            st.info("   다른 인코딩으로 시도해보세요.")
            return None, None
        except Exception as e:
            st.error(f"❌ 파일 텍스트 추출 오류: {str(e)}")
            st.info(f"   오류 타입: {type(e).__name__}")
            return None, None
    
    def analyze_rfp_with_gpt(self, rfp_content: str) -> Dict[str, Any]:
        """GPT-4를 사용하여 RFP 내용 분석"""
        prompt = f"""
        RFP 내용을 11개 카테고리별로 분석하세요:
        
        **요구사항 추출 지침 (매우 중요):**
        1. RFP 문서 전체를 처음부터 끝까지 꼼꼼히 검토하여 모든 요구사항을 찾아서 추출하세요
        2. 요구사항 고유번호 패턴: ECR-XXX-XXX-XX, REQ-XXX-XXX, RFP-XXX-XXX, REQ-XX-XX-XX, REQ-001, REQ-002 등
        3. 각 요구사항별로 분류, 명칭, 세부내용, 산출정보 포함
        4. 문서에 있는 모든 요구사항을 빠짐없이 추출 (고유번호 없는 경우 REQ-GEN-001, REQ-GEN-002 형태로 부여)
        5. 요구사항을 찾을 때는 "요구사항", "기능요구", "비기능요구", "성능요구", "보안요구", "웹접근성", "개선", "구축", "개발" 등의 키워드로 검색하세요
        6. 문서에서 요구사항과 관련된 모든 문장, 문단을 찾아서 각각을 별도의 요구사항으로 추출하세요
        7. 요구사항 상세목록 배열에는 실제 문서에서 찾은 모든 요구사항을 포함하세요 (최소 50개 이상)
        8. 예시를 그대로 복사하지 말고, 실제 RFP 문서에서 찾은 요구사항으로 교체하세요
        
        **요구사항 추출 방법:**
        - 문서의 모든 페이지를 검토하세요
        - 요구사항과 관련된 모든 문장을 찾으세요
        - 각 요구사항을 별도의 항목으로 분리하세요
        - 고유번호가 없는 요구사항도 REQ-GEN-001, REQ-GEN-002 형태로 부여하세요
        - 최소 50개 이상의 요구사항을 추출하세요
        
        **중요: 요구사항 상세목록에는 실제 RFP 문서에서 찾은 모든 요구사항을 포함하세요. 예시를 복사하지 마세요.**
        
        **요구사항 상세목록 형식:**
        각 요구사항은 다음 형식으로 작성하세요:
        {{"요구사항_고유번호": "실제_고유번호", "요구사항_분류": "실제_분류", "요구사항_명칭": "실제_명칭", "요구사항_세부내용": "실제_세부내용", "산출정보": ["실제_산출물"]}}
        
        **요구사항 상세목록 작성 규칙:**
        - 실제 RFP 문서에서 찾은 모든 요구사항을 포함하세요
        - 예시를 복사하지 마세요
        - 최소 50개 이상의 요구사항을 추출하세요
        - 각 요구사항은 별도의 객체로 작성하세요
        
        **요구사항 상세목록 예시 (참고용):**
        [
            {{"요구사항_고유번호": "REQ-001", "요구사항_분류": "기능요구", "요구사항_명칭": "웹접근성 개선", "요구사항_세부내용": "실제 세부내용", "산출정보": ["실제 산출물"]}},
            {{"요구사항_고유번호": "ECR-002", "요구사항_분류": "비기능요구", "요구사항_명칭": "성능 최적화", "요구사항_세부내용": "실제 세부내용", "산출정보": ["실제 산출물"]}},
            {{"요구사항_고유번호": "DAR-003", "요구사항_분류": "보안요구", "요구사항_명칭": "보안 요구사항", "요구사항_세부내용": "실제 세부내용", "산출정보": ["실제 산출물"]}},
            {{"요구사항_고유번호": "FUN-004", "요구사항_분류": "호환성표준", "요구사항_명칭": "호환성 요구사항", "요구사항_세부내용": "실제 세부내용", "산출정보": ["실제 산출물"]}}
        ]
        
        **위 예시는 참고용이며, 실제 RFP 문서에서 찾은 모든 요구사항으로 교체하세요.**
        
        **요구사항 상세목록 작성 시 반드시 지켜야 할 사항:**
        1. 실제 RFP 문서에서 찾은 모든 요구사항을 포함하세요
        2. 예시를 그대로 복사하지 마세요
        3. 최소 50개 이상의 요구사항을 추출하세요
        4. 각 요구사항은 별도의 객체로 작성하세요
        5. 요구사항 고유번호는 실제 문서에서 찾은 것을 사용하세요
        6. 요구사항 분류는 실제 내용에 맞게 분류하세요
        7. 요구사항 명칭은 실제 내용을 요약하여 작성하세요
        8. 요구사항 세부내용은 실제 내용을 상세히 작성하세요
        9. 산출정보는 실제 산출물을 명시하세요
        
        RFP 내용:
        {rfp_content}

        다음 JSON 형식으로 응답해주세요:
        {{
            "1_핵심개요": {{
                "배경목적": "프로젝트 배경 및 목적",
                "범위": "프로젝트 범위 (포함/제외 사항)",
                "기대성과": "비즈니스 목표 및 효과 지표",
                "용어정의": "주요 용어 및 약어 정의",
                "이해관계자": "발주부서 및 이해관계자"
            }},
            "2_일정마일스톤": {{
                "사업기간": "착수일부터 종료일까지",
                "주요마일스톤": "착수/중간점검/시범/검수 일정",
                "제출물일정": "요구서/설계/결과보고 등 제출물 일정",
                "질의응답마감": "Q&A 및 제안서 접수 마감일"
            }},
            "3_예산가격": {{
                "추정예산": "예산 범위 및 상한가",
                "부가세포함": "부가세 포함 여부",
                "가격구성": "라이선스/구축/운영/교육/옵션 비용",
                "지불조건": "선급/중도/준공/검수 연동 지불 조건",
                "원가산출근거": "인력단가, 수량, 산식 등"
            }},
            "4_평가선정기준": {{
                "정량정성배점": "기술/가격 비율 및 배점표",
                "가점감점요건": "레퍼런스, 인증, 현장실사 등",
                "탈락필수요건": "필수 서류 및 자격 미충족 시 탈락 조건"
            }},
            "5_요구사항": {{
                "요구사항_상세목록": [
                    {{"요구사항_고유번호": "REQ-001", "요구사항_분류": "기능요구", "요구사항_명칭": "웹접근성 개선", "요구사항_세부내용": "실제 세부내용", "산출정보": ["실제 산출물"]}},
                    {{"요구사항_고유번호": "ECR-002", "요구사항_분류": "비기능요구", "요구사항_명칭": "성능 최적화", "요구사항_세부내용": "실제 세부내용", "산출정보": ["실제 산출물"]}},
                    {{"요구사항_고유번호": "DAR-003", "요구사항_분류": "보안요구", "요구사항_명칭": "보안 요구사항", "요구사항_세부내용": "실제 세부내용", "산출정보": ["실제 산출물"]}},
                    {{"요구사항_고유번호": "FUN-004", "요구사항_분류": "호환성표준", "요구사항_명칭": "호환성 요구사항", "요구사항_세부내용": "실제 세부내용", "산출정보": ["실제 산출물"]}}
                ],
                "기능요구": "요구사항 고유번호별 핵심 기능 요구사항",
                "인터페이스연계": "시스템 목록, 연계 방식, 주기",
                "데이터": "요구사항 고유번호별 데이터 관련 요구사항",
                "비기능요구": "요구사항 고유번호별 성능, 가용성, 확장성, 보안, 접근성 요구사항",
                "호환성표준": "국가표준, 오픈API, 브라우저/OS 호환성"
            }},
            "6_보안준법": {{
                "인증권한감사": "로그, 분리, 추적성",
                "개인정보컴플라이언스": "ISO27001, ISMS, GDPR 등",
                "망구성암호화": "망구성, 암호화, 키관리",
                "취약점진단": "취약점 진단 및 보안점검 대응"
            }},
            "7_서비스수준운영": {{
                "SLA": "가용성, 응답/복구 시간, 페널티",
                "장애변경배포": "ITSM, CAB 프로세스",
                "모니터링리포팅": "KPI, 주기, 포맷",
                "헬프데스크": "지원 시간 및 티어",
                "교육매뉴얼": "교육, 매뉴얼, 전환운영, 케어기간"
            }},
            "8_품질검수인수": {{
                "산출물목록": "산출물 목록 및 템플릿",
                "테스트계획": "단위/통합/성능/UAT 테스트 계획",
                "인수기준": "인수 기준, 결함 허용치, 재검수 규칙",
                "파일럿PoC": "파일럿/PoC 조건"
            }},
            "9_계약법무": {{
                "계약유형": "총액/단가/성과형 계약",
                "지적재산권": "소스코드 소유 및 사용권",
                "비밀유지": "NDA, 자료반환 조건",
                "손해배상": "손해배상, 지체상금, 보증, 보험",
                "하자보수": "하자보수 기간 및 범위"
            }},
            "10_공급사자격역량": {{
                "참여제한": "업종, 등급, 실적 등 참여 제한",
                "필수자격": "필수 자격 요건",
                "투입인력": "등급, 자격증, 상주 여부",
                "레퍼런스": "유사 프로젝트 규모, 기간, 기술스택"
            }},
            "11_제출형식지시": {{
                "제안서형식": "제안서 형식, 분량, 언어, 파일 규격",
                "필수첨부": "서약서, 인증서, 재무제표 등",
                "제출채널": "제출 채널, 원본/사본 매수",
                "프레젠테이션": "데모/샘플/시연 요구 기준"
            }},
            "기술솔루션매핑": {{"요구사항": "구체적인 기술 솔루션 명"}},
            "핵심키워드": ["핵심 키워드들"]
        }}
        """
        
        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "당신은 RFP 분석 전문가입니다. 주어진 RFP 문서를 정확하게 분석하여 구조화된 정보를 추출합니다."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=2000
            )
            
            content = response.choices[0].message.content
            # JSON 추출
            json_match = re.search(r'\{.*\}', content, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
            else:
                st.error("❌ GPT 응답에서 JSON을 찾을 수 없습니다.")
                return {}
                
        except Exception as e:
            st.error(f"❌ RFP 분석 오류: {str(e)}")
            return {}
    
    def store_rfp_in_search(self, rfp_analysis: Dict[str, Any], rfp_content: str, pdf_title: str = None):
        """분석된 RFP를 Azure AI Search에 저장"""
        try:
            # 임베딩 벡터 생성
            content_vector = self.get_embedding(rfp_content)
            
            # PDF 제목 추출 (우선순위: 전달받은 pdf_title > 프로젝트 개요 > 기본값)
            if pdf_title:
                document_title = pdf_title
            else:
                project_overview = ""
                if "1_핵심개요" in rfp_analysis:
                    core_overview = rfp_analysis["1_핵심개요"]
                    project_overview = f"{core_overview.get('배경목적', '')} {core_overview.get('기대성과', '')}".strip()
                document_title = project_overview[:100] if project_overview else "RFP Document"
            
            # 예산 범위 추출
            budget_range = ""
            if "3_예산가격" in rfp_analysis:
                budget_info = rfp_analysis["3_예산가격"]
                budget_range = budget_info.get("추정예산", "")
            
            # 제출 마감일 추출
            submission_deadline = ""
            if "2_일정마일스톤" in rfp_analysis:
                schedule_info = rfp_analysis["2_일정마일스톤"]
                submission_deadline = schedule_info.get("질의응답마감", "")
            
            # 기술 요구사항 추출
            technical_requirements = []
            if "5_요구사항" in rfp_analysis:
                requirements_info = rfp_analysis["5_요구사항"]
                technical_requirements = [
                    requirements_info.get("기능요구", ""),
                    requirements_info.get("비기능요구", ""),
                    requirements_info.get("인터페이스연계", ""),
                    requirements_info.get("데이터", ""),
                    requirements_info.get("호환성표준", "")
                ]
                technical_requirements = [req for req in technical_requirements if req]
            
            # 평가 기준 추출
            evaluation_criteria = {}
            if "4_평가선정기준" in rfp_analysis:
                eval_info = rfp_analysis["4_평가선정기준"]
                evaluation_criteria = {
                    "정량정성배점": eval_info.get("정량정성배점", ""),
                    "가점감점요건": eval_info.get("가점감점요건", ""),
                    "탈락필수요건": eval_info.get("탈락필수요건", "")
                }
            
            document = {
                "id": f"rfp_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                "title": document_title,
                "content": rfp_content,
                "requirements": json.dumps(technical_requirements, ensure_ascii=False),
                "project_type": rfp_analysis.get("project_type", ""),
                "budget_range": budget_range,
                "submission_deadline": submission_deadline,
                "evaluation_criteria": json.dumps(evaluation_criteria, ensure_ascii=False),
                "created_date": datetime.now().isoformat() + "Z",
                "content_vector": content_vector
            }
            
            result = self.search_client.upload_documents([document])
            return len(result) > 0
            
        except Exception as e:
            st.error(f"❌ RFP 저장 오류: {str(e)}")
            return False
    
    def search_similar_rfps(self, current_rfp_keywords: List[str], limit: int = 5):
        """유사한 RFP 검색"""
        try:
            # 키워드를 조합하여 검색 쿼리 생성
            search_query = " ".join(current_rfp_keywords)
            
            # 벡터 검색을 위한 임베딩 생성
            query_vector = self.get_embedding(search_query)
            
            if not query_vector:
                return []
            
            # 벡터 검색 수행
            vector_query = VectorizedQuery(vector=query_vector, k_nearest_neighbors=limit, fields="content_vector")
            
            results = self.search_client.search(
                search_text=search_query,
                vector_queries=[vector_query],
                select=["title", "project_type", "requirements", "evaluation_criteria", "created_date"],
                top=limit
            )
            
            similar_rfps = []
            for result in results:
                similar_rfps.append({
                    "title": result["title"],
                    "project_type": result["project_type"],
                    "requirements": result["requirements"],
                    "evaluation_criteria": result["evaluation_criteria"],
                    "created_date": result["created_date"],
                    "score": result.get("@search.score", 0)
                })
            
            return similar_rfps
            
        except Exception as e:
            st.error(f"❌ 유사 RFP 검색 오류: {str(e)}")
            return []
    
    def ask_question_about_rfp(self, question: str, rfp_content: str, rfp_analysis: Dict[str, Any] = None) -> str:
        """RFP 내용에 대한 질의응답"""
        try:
            # RFP 분석 결과가 있으면 구조화된 정보도 포함
            analysis_context = ""
            if rfp_analysis:
                analysis_context = f"""
                
                RFP 분석 결과:
                {json.dumps(rfp_analysis, ensure_ascii=False, indent=2)}
                """
            
            prompt = f"""
            RFP 문서 내용을 바탕으로 질문에 답변해주세요.
            
            RFP 내용:
            {rfp_content[:3000]}  # 토큰 제한을 고려하여 3000자로 제한
            
            {analysis_context}
            
            사용자 질문: {question}
            
            답변 지침:
            1. RFP 문서에서 직접 찾을 수 있는 정보를 우선 제공
            2. 구체적인 수치, 조건, 일정 등 정확히 명시
            3. 관련 조항이나 섹션 참조
            4. 정보가 명확하지 않으면 "문서에서 명확한 정보를 찾을 수 없습니다" 표시
            5. 한국어로 간결하게 답변
            
            답변:
            """
            
            response = self.openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "당신은 RFP 문서 분석 전문가입니다. RFP 문서의 내용을 정확하게 분석하여 사용자의 질문에 구체적이고 정확한 답변을 제공합니다."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=1000
            )
            
            return response.choices[0].message.content
            
        except Exception as e:
            st.error(f"❌ 질의응답 처리 오류: {str(e)}")
            return "질문 처리 중 오류가 발생했습니다. 다시 시도해주세요."

    # def generate_proposal_draft(self, rfp_analysis: Dict[str, Any], similar_rfps: List[Dict]) -> str:
    #     """제안서 초안 생성"""
    #     similar_rfps_text = ""
    #     for i, rfp in enumerate(similar_rfps[:3], 1):
    #         similar_rfps_text += f"\n{i}. {rfp['title']} (유형: {rfp['project_type']})\n"
    #     
    #     # RFP 분석 정보를 구조화된 형태로 정리
    #     rfp_summary = ""
    #     for category_key, category_data in rfp_analysis.items():
    #         if category_key.startswith(('1_', '2_', '3_', '4_', '5_', '6_', '7_', '8_', '9_', '10_', '11_')):
    #             category_name = category_key.replace('_', ' ').replace('1 ', '1. ').replace('2 ', '2. ').replace('3 ', '3. ').replace('4 ', '4. ').replace('5 ', '5. ').replace('6 ', '6. ').replace('7 ', '7. ').replace('8 ', '8. ').replace('9 ', '9. ').replace('10 ', '10. ').replace('11 ', '11. ')
    #             rfp_summary += f"\n{category_name}:\n"
    #             
    #             # 요구사항의 경우 상세목록이 있으면 고유번호별로 표시
    #             if category_key == "5_요구사항" and isinstance(category_data, dict) and "요구사항_상세목록" in category_data:
    #                 requirements_list = category_data.get("요구사항_상세목록", [])
    #                 if requirements_list:
    #                     rfp_summary += "  요구사항 상세목록:\n"
    #                     for req in requirements_list:
    #                         rfp_summary += f"    - {req.get('요구사항_고유번호', 'N/A')}: {req.get('요구사항_명칭', 'N/A')}\n"
    #                         rfp_summary += f"      분류: {req.get('요구사항_분류', 'N/A')}\n"
    #                         rfp_summary += f"      세부내용: {req.get('요구사항_세부내용', 'N/A')}\n"
    #                         if req.get('산출정보'):
    #                             rfp_summary += f"      산출정보: {', '.join(req.get('산출정보', []))}\n"
    #                         rfp_summary += "\n"
    #             
    #             if isinstance(category_data, dict):
    #                 for key, value in category_data.items():
    #                     if value and key != "요구사항_상세목록":
    #                         rfp_summary += f"  - {key}: {value}\n"
    #             else:
    #                 rfp_summary += f"  {category_data}\n"
    #     
    #     prompt = f"""
    #     다음 RFP 분석 결과를 바탕으로 전문적인 제안서 초안을 작성해주세요:

    #     RFP 분석 결과:
    #     {rfp_summary}

    #     기술 솔루션 매핑: {json.dumps(rfp_analysis.get('기술솔루션매핑', {}), ensure_ascii=False)}

    #     유사 프로젝트 사례:{similar_rfps_text}

    #     **제안서 작성 지침:**
    #     - 요구사항 고유번호(예: ECR-HWR-SVR-02, REQ-XXX-XXX 등)를 명시하여 각 요구사항에 대한 구체적인 솔루션을 제시하세요
    #     - 각 요구사항별로 어떻게 해결할 것인지 명확하게 설명하세요
    #     - 요구사항 고유번호와 함께 산출물도 언급하세요 (예: ECR-HWR-SVR-02 요구사항에 대한 아키텍처 설계서 제출)

    #     다음 구조로 제안서 초안을 작성해주세요:

    #     **Ⅰ. 제안 개요**
    #     - 사업이해도

    #     **Ⅱ. 제안 업체 일반 (KT DS)**
    #     1. 일반 현황
    #     2. 조직 및 인원
    #     3. 주요 사업 내용
    #     4. 주요 사업 실적

    #     **Ⅲ. 프로젝트 수행 부문**
    #     1. 추진 전략 및 개발 방법론
    #     2. 시스템 구성도
    #     3. 시스템 구축
    #         > 3.1 개발대상업무 내역 및 구성요건
    #         > 3.2 연계 범위
    #         > 3.3 표준화 요건
    #         > 3.4 보안 및 웹 표준화, 성능시험 요구사항
    #         > 3.5 테스트 수행 방안

    #     **Ⅳ. 프로젝트 관리 부문**
    #     1. 프로젝트 관리 방법론
    #     2. 추진 일정 계획
    #     3. 투입 인력 및 이력 사항

    #     **Ⅴ. 지원 부문**
    #     1. 교육 훈련 계획
    #     2. 기술지원 계획
    #     3. 하자보수 계획
    #     4. 안정화
    #     """
    #     
    #     try:
    #         response = self.openai_client.chat.completions.create(
    #             model="gpt-4o-mini",
    #             messages=[
    #                 {"role": "system", "content": "당신은 경험이 풍부한 제안서 작성 전문가입니다. 기술적으로 정확하고 설득력 있는 제안서를 작성합니다."},
    #                 {"role": "user", "content": prompt}
    #             ],
    #             temperature=0.5,
    #             max_tokens=3000
    #         )
    #         
    #         return response.choices[0].message.content
    #         
    #     except Exception as e:
    #         st.error(f"❌ 제안서 초안 생성 오류: {str(e)}")
    #         return ""
